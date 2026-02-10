from __future__ import annotations

import json
import subprocess
from pathlib import Path

from docx import Document


def build_service_description(items: list[dict]) -> str:
    rows: list[str] = []
    for idx, item in enumerate(items, start=1):
        note_text = "; ".join(f"{k}: {v}" for k, v in item["notes"].items() if v)
        if note_text:
            rows.append(f"{idx}. {item['service_name']} ({item['qty']} x {item['price']:.2f}) — {note_text}")
        else:
            rows.append(f"{idx}. {item['service_name']} ({item['qty']} x {item['price']:.2f})")
    return "\n".join(rows)


def _replace_in_paragraph(paragraph, context: dict[str, str]) -> None:
    if not paragraph.text:
        return
    full_text = paragraph.text
    for key, value in context.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
    if full_text != paragraph.text:
        paragraph.clear()
        paragraph.add_run(full_text)


def replace_docx_variables(template_path: Path, output_path: Path, context: dict[str, str]) -> None:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, context)

    doc.save(output_path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    """Convert docx into pdf via docx2pdf or LibreOffice. Returns status and details."""
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return True, "docx2pdf"
    except Exception:
        pass

    libreoffice_bin = "libreoffice"
    try:
        result = subprocess.run(
            [
                libreoffice_bin,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        generated_pdf = docx_path.with_suffix(".pdf")
        if generated_pdf.exists() and generated_pdf != pdf_path:
            generated_pdf.replace(pdf_path)
        return True, result.stdout.strip() or "libreoffice"
    except Exception as exc:
        pdf_path.write_text(
            "PDF conversion is unavailable in current environment."
            f"\nSource DOCX: {docx_path.name}\nError: {exc}\n",
            encoding="utf-8",
        )
        return False, str(exc)


def flatten_notes(items: list[dict]) -> dict[str, str]:
    flat: dict[str, str] = {}
    for idx, item in enumerate(items, start=1):
        for key, value in item["notes"].items():
            flat[f"item_{idx}_{key}"] = value
    return flat


def items_to_json(items: list[dict]) -> str:
    return json.dumps(items, ensure_ascii=False)
